import streamlit as st
import os
import sys
import tempfile
import uuid
import sqlite3
from typing import List, Dict, Any
from dotenv import load_dotenv
from PyPDF2 import PdfReader
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from langchain_community.vectorstores import Chroma, FAISS
from langchain_openai import ChatOpenAI
from langchain.memory import ConversationBufferMemory
from langchain.chains import ConversationalRetrievalChain
from langchain_community.callbacks import get_openai_callback
from langchain.docstore.document import Document as LangchainDocument
from langchain.prompts import PromptTemplate
from deep_translator import GoogleTranslator

def check_openai_key():
    """Check if OpenAI API key is set"""
    api_key = st.secrets.get("OPENAI_API_KEY") or os.getenv('OPENAI_API_KEY')
    if not api_key:
        st.error("⚠️ OpenAI API key not found. Please set it in your Streamlit secrets or .env file!")
        return False
    return True

def check_sqlite_version():
    """Check if SQLite version meets ChromaDB requirements"""
    sqlite_version = sqlite3.sqlite_version_info
    required_version = (3, 35, 0)
    return sqlite_version >= required_version

def translate_text_batch(text: str, batch_size: int = 1000) -> str:
    """Translate text in batches while preserving context"""
    try:
        # Split text into sentences to preserve context
        sentences = text.split('.')
        batches = []
        current_batch = []
        current_length = 0
        
        for sentence in sentences:
            sentence = sentence.strip() + '.'
            if current_length + len(sentence) > batch_size:
                if current_batch:
                    batches.append(' '.join(current_batch))
                current_batch = [sentence]
                current_length = len(sentence)
            else:
                current_batch.append(sentence)
                current_length += len(sentence)
        
        if current_batch:
            batches.append(' '.join(current_batch))
        
        translated_batches = []
        for batch in batches:
            translated = GoogleTranslator(source='auto', target='en').translate(batch)
            translated_batches.append(translated)
        
        return ' '.join(translated_batches)
    except Exception as e:
        st.error(f"Translation error: {str(e)}")
        return text

def get_docx_text(file) -> LangchainDocument:
    """Extract text from a .docx file"""
    doc = Document(file)
    text = []
    for paragraph in doc.paragraphs:
        text.append(paragraph.text)
    return LangchainDocument(
        page_content='\n'.join(text),
        metadata={"source": file.name, "type": "docx"}
    )

def get_pdf_text(file) -> LangchainDocument:
    """Extract text from a PDF file"""
    pdf_reader = PdfReader(file)
    text = []
    for page in pdf_reader.pages:
        text.append(page.extract_text())
    return LangchainDocument(
        page_content='\n'.join(text),
        metadata={"source": file.name, "type": "pdf", "pages": len(pdf_reader.pages)}
    )

def process_documents(docs) -> List[LangchainDocument]:
    """Process multiple documents and return a list of Document objects"""
    documents = []
    for doc in docs:
        if doc.name.lower().endswith('.pdf'):
            documents.append(get_pdf_text(doc))
        elif doc.name.lower().endswith('.docx'):
            documents.append(get_docx_text(doc))
    return documents

def get_text_chunks(documents: List[LangchainDocument]) -> List[LangchainDocument]:
    """Split documents into chunks while preserving metadata"""
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        length_function=len,
        add_start_index=True,
    )
    
    chunks = []
    for doc in documents:
        # Translate content if needed
        translated_text = translate_text_batch(doc.page_content)
        doc.page_content = translated_text
        doc_chunks = text_splitter.split_documents([doc])
        chunks.extend(doc_chunks)
    
    return chunks

def get_vectorstore(text_chunks: List[LangchainDocument]):
    """Create vectorstore from document chunks"""
    embeddings = OpenAIEmbeddings()
    
    # Check SQLite version and use appropriate vector store
    if check_sqlite_version():
        st.session_state.debug_info.append("Using ChromaDB vector store")
        # Use in-memory ChromaDB without persistence and with reduced dimension
        vectorstore = Chroma.from_documents(
            documents=text_chunks,
            embedding=embeddings,
            collection_metadata={"hnsw:space": "cosine", "hnsw:construction_ef": 80, "hnsw:M": 8}
        )
    else:
        st.session_state.debug_info.append("Using FAISS vector store (SQLite version < 3.35.0)")
        # Use FAISS as fallback
        vectorstore = FAISS.from_documents(
            documents=text_chunks,
            embedding=embeddings
        )
    return vectorstore

def get_conversation_chain(vectorstore) -> ConversationalRetrievalChain:
    """Create conversation chain with custom prompt and retrieval settings"""
    llm = ChatOpenAI(
        temperature=0.0,
        model_name="gpt-3.5-turbo",
        max_tokens=500
    )
    
    # Simplified prompt template
    prompt_template = """Answer based on the context below. If you're unsure, say "I don't have enough information."

Context: {context}

Chat History: {chat_history}

Question: {question}

Answer:"""
    
    PROMPT = PromptTemplate(
        input_variables=["context", "chat_history", "question"],
        template=prompt_template
    )
    
    memory = ConversationBufferMemory(
        memory_key='chat_history',
        return_messages=True,
        output_key='answer',
        max_token_limit=1000
    )
    
    conversation_chain = ConversationalRetrievalChain.from_llm(
        llm=llm,
        retriever=vectorstore.as_retriever(
            search_kwargs={"k": 5}
        ),
        memory=memory,
        combine_docs_chain_kwargs={"prompt": PROMPT},
        return_source_documents=False
    )
    
    return conversation_chain

def main():
    # Load environment variables
    load_dotenv()
    
    # Initialize session state
    if 'conversation' not in st.session_state:
        st.session_state.conversation = None
    if 'chat_history' not in st.session_state:
        st.session_state.chat_history = None
    if 'debug_info' not in st.session_state:
        st.session_state.debug_info = []
    if 'processed_docs' not in st.session_state:
        st.session_state.processed_docs = []
    if 'vectorstore' not in st.session_state:
        st.session_state.vectorstore = None

    try:
        st.set_page_config(
            page_title="Chat with Documents",
            page_icon="💬",
            layout="wide",
            initial_sidebar_state="expanded"
        )
        st.header("Chat with your Documents 💬")
        
        # Add debug info section in sidebar
        with st.sidebar:
            st.subheader("Debug Information")
            st.text(f"Python version: {sys.version}")
            st.text(f"Working directory: {os.getcwd()}")
            if st.session_state.debug_info:
                for info in st.session_state.debug_info:
                    st.text(info)
            
            # Clear memory button
            if st.button("Clear Memory"):
                st.session_state.conversation = None
                st.session_state.chat_history = None
                st.session_state.vectorstore = None
                st.session_state.debug_info = []
                st.success("Memory cleared!")
            
            st.subheader("Your documents")
            uploaded_files = st.file_uploader(
                "Upload your PDFs/DOCXs here",
                accept_multiple_files=True,
                type=['pdf', 'docx']
            )

            if uploaded_files:
                st.session_state.debug_info = []  # Clear previous debug info
                
                # Process documents
                with st.spinner("Processing documents..."):
                    try:
                        # Clear previous conversation when new documents are uploaded
                        st.session_state.conversation = None
                        st.session_state.chat_history = None
                        
                        documents = process_documents(uploaded_files)
                        st.session_state.debug_info.append(f"Processed {len(documents)} documents")
                        
                        text_chunks = get_text_chunks(documents)
                        st.session_state.debug_info.append(f"Created {len(text_chunks)} text chunks")
                        
                        vectorstore = get_vectorstore(text_chunks)
                        st.session_state.vectorstore = vectorstore
                        st.session_state.debug_info.append("Created vector store")
                        
                        # Create conversation chain
                        st.session_state.conversation = get_conversation_chain(vectorstore)
                        st.session_state.debug_info.append("Created conversation chain")
                        
                        st.success("Documents processed successfully!")
                    except Exception as e:
                        st.error(f"Error processing documents: {str(e)}")
                        st.session_state.debug_info.append(f"Error: {str(e)}")
                        return

        # Chat interface
        if st.session_state.conversation is not None:
            user_question = st.text_input("Ask a question about your documents:")
            
            if user_question:
                with st.spinner("Thinking..."):
                    with get_openai_callback() as cb:
                        try:
                            response = st.session_state.conversation({
                                "question": user_question
                            })
                            st.write(response["answer"])
                            
                            # Display token usage
                            st.sidebar.text(f"Tokens used: {cb.total_tokens}")
                            st.sidebar.text(f"Cost: ${cb.total_cost:.4f}")
                        except Exception as e:
                            st.error(f"Error getting response: {str(e)}")
                            st.session_state.debug_info.append(f"Error: {str(e)}")
        else:
            st.info("Upload documents to start chatting!")

    except Exception as e:
        st.error(f"Application error: {str(e)}")
        st.session_state.debug_info.append(f"Fatal error: {str(e)}")

if __name__ == '__main__':
    main()
